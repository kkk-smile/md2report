"""
中间结构 → docx 生成器
依赖 report_utils.py 的格式函数
"""

import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from report_utils import (
    new_doc, build_cover_personal, build_cover_group,
    FONT_HEITI, FONT_SONGTI, FONT_TNR,
    _set_run_font, _set_spacing,
)

# ── 默认格式（与 report_utils 保持一致）────────────────────────────
DEFAULT_FMT = {
    "school":       "莆田学院新工科产业学院",
    "h1_font":      "黑体",
    "h1_size":      16,
    "h1_bold":      False,
    "h2_font":      "宋体",
    "h2_size":      16,
    "h2_bold":      True,
    "h3_font":      "宋体",
    "h3_size":      14,
    "h3_bold":      True,
    "body_font":    "宋体",
    "body_size":    12,
    "line_spacing": 1.5,
    "margin_top":    2.54,
    "margin_bottom": 2.54,
    "margin_left":   3.17,
    "margin_right":  2.54,
}


def _merge_fmt(user_fmt: dict) -> dict:
    """用户传入的格式参数覆盖默认值"""
    fmt = DEFAULT_FMT.copy()
    if user_fmt:
        for k, v in user_fmt.items():
            if k in fmt and v not in (None, "", ):
                # 数值类型做类型转换
                if isinstance(fmt[k], float):
                    try: fmt[k] = float(v)
                    except: pass
                elif isinstance(fmt[k], int):
                    try: fmt[k] = int(v)
                    except: pass
                elif isinstance(fmt[k], bool):
                    fmt[k] = bool(v)
                else:
                    fmt[k] = v
    return fmt


def _add_h1(doc, text, fmt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p.paragraph_format, space_before=12, space_after=6,
                 multiple=fmt["line_spacing"])
    _set_run_font(p.add_run(text), fmt["h1_font"], FONT_TNR,
                  fmt["h1_size"], bold=fmt["h1_bold"])
    return p


def _add_h2(doc, text, fmt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p.paragraph_format, space_before=6, space_after=3,
                 multiple=fmt["line_spacing"])
    _set_run_font(p.add_run(text), fmt["h2_font"], FONT_TNR,
                  fmt["h2_size"], bold=fmt["h2_bold"])
    return p


def _add_h3(doc, text, fmt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p.paragraph_format, space_before=4, space_after=2,
                 multiple=fmt["line_spacing"])
    _set_run_font(p.add_run(text), fmt["h3_font"], FONT_TNR,
                  fmt["h3_size"], bold=fmt["h3_bold"])
    return p


def _add_body(doc, text, fmt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    _set_spacing(pf, multiple=fmt["line_spacing"])
    pf.first_line_indent = Pt(fmt["body_size"] * 2)
    _set_run_font(p.add_run(text), fmt["body_font"], FONT_TNR, fmt["body_size"])
    return p


def _add_code(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent  = Pt(24)
    pf.space_before = Pt(1)
    pf.space_after  = Pt(1)
    pf.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
    return p


def build_docx(cover: dict, blocks: list, fmt: dict = None) -> bytes:
    """
    cover 字段：
      type      : "personal" | "group"
      name/sid  : 个人信息
      members   : [(姓名, 学号), ...]（小组）
      course/topic/major/teacher : 封面信息

    fmt 字段（均可选，不传则用默认值）：
      school, h1_font, h1_size, h1_bold,
      h2_font, h2_size, h2_bold,
      h3_font, h3_size, h3_bold,
      body_font, body_size, line_spacing,
      margin_top, margin_bottom, margin_left, margin_right

    返回 docx 字节流
    """
    f = _merge_fmt(fmt)

    doc = new_doc()
    # 应用自定义页边距
    for sec in doc.sections:
        sec.top_margin    = Cm(f["margin_top"])
        sec.bottom_margin = Cm(f["margin_bottom"])
        sec.left_margin   = Cm(f["margin_left"])
        sec.right_margin  = Cm(f["margin_right"])

    # 封面（学校名从 fmt 里取）
    school = f["school"]
    if cover.get("type") == "group":
        build_cover_group(
            doc,
            members=cover.get("members", []),
            course=cover.get("course", ""),
            topic=cover.get("topic", ""),
            major=cover.get("major", ""),
            teacher=cover.get("teacher", ""),
            school=school,
        )
    else:
        build_cover_personal(
            doc,
            name=cover.get("name", ""),
            sid=cover.get("sid", ""),
            course=cover.get("course", ""),
            topic=cover.get("topic", ""),
            major=cover.get("major", ""),
            teacher=cover.get("teacher", ""),
            school=school,
        )

    # 正文
    for block in blocks:
        t = block.get("type")
        if t == "h1":
            _add_h1(doc, block["text"], f)
        elif t == "h2":
            _add_h2(doc, block["text"], f)
        elif t == "h3":
            _add_h3(doc, block["text"], f)
        elif t == "body":
            _add_body(doc, block["text"], f)
        elif t == "code":
            for line in block.get("lines", []):
                _add_code(doc, line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
