"""
课程设计报告通用格式工具库
莆田学院新工科产业学院 - 大数据分析与可视化项目实战

格式规范：
  封面：楷体36（学校名）/ 黑体42（课程设计）/ 宋体22（信息行）
  一级标题：黑体，三号(16pt)，段前12pt，段后6pt，1.5倍行距
  二级标题：宋体 + Times New Roman，三号(16pt)，加粗，1.5倍行距
  三级标题：宋体 + Times New Roman，四号(14pt)，加粗，1.5倍行距
  正文：宋体 + Times New Roman，小四(12pt)，1.5倍行距，首行缩进2字符
  代码块：Courier New，10pt，左缩进
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

FONT_KAITI  = '楷体'
FONT_HEITI  = '黑体'
FONT_SONGTI = '宋体'
FONT_TNR    = 'Times New Roman'

SIZE_COVER_SCHOOL = 36
SIZE_COVER_TITLE  = 42
SIZE_COVER_INFO   = 22
SIZE_COVER_DATE   = 18
SIZE_H1           = 16
SIZE_H2           = 16
SIZE_H3           = 14
SIZE_BODY         = 12
SIZE_CODE         = 10

MARGIN_TOP    = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)
MARGIN_LEFT   = Cm(3.17)
MARGIN_RIGHT  = Cm(2.54)


def _set_run_font(run, cn_font, en_font, size_pt, bold=False):
    run.font.name = en_font
    run.font.size = Pt(size_pt)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn('w:eastAsia'), cn_font)
    rpr.rFonts.set(qn('w:ascii'),    en_font)
    rpr.rFonts.set(qn('w:hAnsi'),    en_font)


def _set_spacing(pf, multiple=1.5, space_before=0, space_after=0):
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = multiple
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)


def new_doc():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = MARGIN_TOP
        sec.bottom_margin = MARGIN_BOTTOM
        sec.left_margin   = MARGIN_LEFT
        sec.right_margin  = MARGIN_RIGHT
    return doc


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p.paragraph_format, space_before=12, space_after=6)
    _set_run_font(p.add_run(text), FONT_HEITI, FONT_TNR, SIZE_H1, bold=False)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p.paragraph_format, space_before=6, space_after=3)
    _set_run_font(p.add_run(text), FONT_SONGTI, FONT_TNR, SIZE_H2, bold=True)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p.paragraph_format, space_before=4, space_after=2)
    _set_run_font(p.add_run(text), FONT_SONGTI, FONT_TNR, SIZE_H3, bold=True)
    return p


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    _set_spacing(pf)
    if indent:
        pf.first_line_indent = Pt(SIZE_BODY * 2)
    _set_run_font(p.add_run(text), FONT_SONGTI, FONT_TNR, SIZE_BODY)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent  = Pt(24)
    pf.space_before = Pt(1)
    pf.space_after  = Pt(1)
    pf.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(SIZE_CODE)
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
    return p


def add_blank(doc):
    p = doc.add_paragraph()
    _set_spacing(p.paragraph_format)
    _set_run_font(p.add_run(''), FONT_SONGTI, FONT_TNR, SIZE_BODY)
    return p


def add_page_break(doc):
    doc.add_page_break()


def _cover_line(doc, text, cn_font, size_pt, bold=False, space_before=4, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name = cn_font
    run.font.size = Pt(size_pt)
    run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), cn_font)
    return p


def build_cover_personal(doc, name, sid,
                         course='大数据分析与可视化项目实战',
                         topic='', major='大数据技术2301',
                         teacher='（指导教师姓名）',
                         school='莆田学院新工科产业学院'):
    _cover_line(doc, '', FONT_SONGTI, 12, space_before=60, space_after=0)
    _cover_line(doc, school, FONT_KAITI, SIZE_COVER_SCHOOL, space_before=12, space_after=12)
    _cover_line(doc, '', FONT_SONGTI, 12, space_before=6, space_after=0)
    _cover_line(doc, '课 程 设 计', FONT_HEITI, SIZE_COVER_TITLE, space_before=6, space_after=24)
    _cover_line(doc, '', FONT_SONGTI, 12, space_before=12, space_after=0)

    for label, value in [
        ('课程名称：', course), ('题    目：', topic),
        ('学生姓名：', name),   ('学    号：', sid),
        ('专业班级：', major),  ('指导教师：', teacher),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(15.6)
        _set_run_font(p.add_run(label + value), FONT_SONGTI, FONT_TNR, SIZE_COVER_INFO)

    _cover_line(doc, '', FONT_SONGTI, 12, space_before=24, space_after=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    _set_run_font(p.add_run('202   年     月     日'), FONT_SONGTI, FONT_TNR, SIZE_COVER_DATE, bold=True)
    add_page_break(doc)


def build_cover_group(doc, members,
                      course='大数据分析与可视化项目实战',
                      topic='', major='大数据技术2301',
                      teacher='（指导教师姓名）',
                      school='莆田学院新工科产业学院'):
    _cover_line(doc, '', FONT_SONGTI, 12, space_before=60, space_after=0)
    _cover_line(doc, school, FONT_KAITI, SIZE_COVER_SCHOOL, space_before=12, space_after=12)
    _cover_line(doc, '', FONT_SONGTI, 12, space_before=6, space_after=0)
    _cover_line(doc, '课 程 设 计', FONT_HEITI, SIZE_COVER_TITLE, space_before=6, space_after=24)
    _cover_line(doc, '', FONT_SONGTI, 12, space_before=12, space_after=0)

    def cline(text, space_after=15.6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)
        _set_run_font(p.add_run(text), FONT_SONGTI, FONT_TNR, SIZE_COVER_INFO)

    cline(f'课程名称：{course}')
    cline(f'题    目：{topic}')
    for name, sid in members:
        cline(f'姓名学号：{name}  {sid}', space_after=7.8)
    cline(f'专业班级：{major}')
    cline(f'指导教师：{teacher}')

    _cover_line(doc, '', FONT_SONGTI, 12, space_before=24, space_after=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    _set_run_font(p.add_run('202   年     月     日'), FONT_SONGTI, FONT_TNR, SIZE_COVER_DATE, bold=True)
    add_page_break(doc)
